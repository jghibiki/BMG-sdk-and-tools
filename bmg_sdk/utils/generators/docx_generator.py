from bmg_sdk.compendium.models.compendium import Compendium
from bmg_sdk.utils.common import Paths, get_character_dir_name

from tqdm import tqdm
from docx import Document
from docx.shared import Inches


class DocxGenerator:
    def __init__(self, compendium: Compendium):
        Paths.create_docx_output_dir()
        self.compendium = compendium

    @staticmethod
    def write_document(document, affiliation, doc_number):
        margin = 0.5
        for section in document.sections:
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
        document.save(Paths.docx_output / f"{affiliation.name}_{doc_number}.docx")

    def generate(self):

        for affiliation in tqdm(self.compendium.affiliations.all):
            print(affiliation)
            document = Document()
            doc_number = 1
            doc_empty = False
            count = 0
            max_cards = 18
            for character in affiliation.all_characters:
                p = document.add_paragraph()
                r = p.add_run()
                card_dir = Paths.card_output / get_character_dir_name(character)
                front = (card_dir / "front.png")
                back = (card_dir / "back.png")
                for side in (front, back):
                    r.add_picture(
                        str(side),
                        width=Inches(3.5)
                    )

                count += 1
                if count >= max_cards:
                    self.write_document(document, affiliation, doc_number)
                    document = Document()
                    doc_number += 1
                    count = 0

            if not doc_empty:
                self.write_document(document, affiliation, doc_number)
