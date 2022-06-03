import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from shutil import rmtree
from pathlib import Path
from PIL import Image
from compendium.compendium_loader import CompendiumLoader
from compendium.models.compendium import Compendium
from tqdm import tqdm
from docx import Document
from docx.shared import Inches

output = Path("./output")
card_output = output / "cards"
sheet_output = output / "sheets"
docx_output = output / "docx"


def setup_browser():
    return webdriver.Firefox()

def get_character_dir_name(character):
    name = get_sanitized_name(character)
    return f"{character.id}_{name}"

def get_sanitized_name(character):
    return (f"{character.alias}__{character.name}"
            .replace(" ", "_")
            .replace('"', "")
            )



def scrape(compendium):
    browser = setup_browser()

    for character in tqdm(compendium.characters.all):
        browser.get(character.get_card_url())
        delay = 5

        try:
            # ensure div exists
            (
                WebDriverWait(browser, delay)
                    .until(
                        EC.visibility_of_element_located(
                            (By.CLASS_NAME, "Card_card__39xAa")
                        )
                    )
            )

            # wait till image loads
            browser.execute_async_script("""
                var done = arguments[0];
                var a = new Image;
                a.onload = () => done(true);
                a.src = document.getElementsByClassName( 'Card_card__39xAa' )[0].style.backgroundImage.replace('url(\"', "").replace('\")', "")
            """)

            # scrape parent element
            elem = browser.find_element(By.CLASS_NAME, "Card_characters__iuvFJ")

            name = get_sanitized_name(character)
            children = elem.find_elements(By.XPATH, "./*")
            front = children[0]
            back = children[1]
            output_dir = card_output / get_character_dir_name(character)
            output_dir.mkdir()
            front.screenshot(str(output_dir / "front.png"))
            back.screenshot(str(output_dir / "back.png"))
            print(f"Scraped {name}")
        except TimeoutException:
            print("Load took too long")
    browser.close()
    exit(0)


def build_sheet(characters, sheet_number):
    # max tts supports
    w = 10
    h = 7

    sample_dir_name = card_output / get_character_dir_name(characters[0])

    (card_w, card_h) = Image.open(sample_dir_name / "front.png").size
    grid_size = (
        card_w * w,
        card_h * h
    )

    for side in ("front", "back"):
        grid = Image.new("RGB", size=grid_size)

        for c in characters:
            card_dir = card_output / get_character_dir_name(c)
            card_img = Image.open(card_dir / f"{side}.png")
            relative_id = c.id - (70 * (sheet_number - 1))
            offset_x = (relative_id % w) * card_w
            offset_y = (relative_id // w) * card_h
            grid.paste(card_img, box=(offset_x, offset_y))

        output_file = sheet_output / f"sheet_{sheet_number}_{side}.png"
        grid.save(output_file)
        print(f"Exported sheet {sheet_number} - {side}")


def build_sheets(compendium):
    sheet_number = 1
    sheet = []

    # assign sheets by id, leaving blanks for spots that don't have a charracter with
    # matching id.
    for character in compendium.characters.all:
        if character.id // 70 != sheet_number - 1:
            build_sheet(sheet, sheet_number)
            sheet_number += 1
            sheet = []

        sheet.append(character)

def generate_docx(compendium: Compendium):

    for affiliation in tqdm(compendium.affiliations.all):
        print(affiliation)
        document = Document()
        for character in affiliation.all_characters:
            p = document.add_paragraph()
            r = p.add_run()
            card_dir = card_output / get_character_dir_name(character)
            front = (card_dir / "front.png")
            back = (card_dir / "back.png")
            for side in (front, back):
                r.add_picture(
                    str(side),
                    width=Inches(3.5)
                )

        margin = 0.5
        for section in document.sections:
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
        document.save(docx_output / f"{affiliation.name}.docx")


def main():
    if output.exists():
        rmtree(output)

    output.mkdir()
    card_output.mkdir()
    sheet_output.mkdir()
    docx_output.mkdir()


    compendium = CompendiumLoader().load()

    print("Scraping Cards")
    scrape(compendium)

    print("Building sheets")
    build_sheets(compendium)

    print("generate docx")
    generate_docx(compendium)


if __name__ == '__main__':
    main()

